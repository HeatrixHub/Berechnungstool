"""
tab3_bericht_logic.py
Überarbeitete Version:
- Einheitliche Plotfarbe wie in Tab1
- Tabelle mit Schicht, Material, Dicke, Temperaturen, k_mittel
- Nur noch eine kompakte Kopfzeile (Projektname fett)
"""

import os
import datetime
from typing import Dict
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches

pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))


def build_report_content(project: Dict, author: str, comment: str = "") -> str:
    """Erzeugt den Vorschautext für den Bericht (ohne doppelte Ergebnisdarstellung)."""
    try:
        if not isinstance(project, dict):
            project = project.__dict__

        text = []
        if author:
            text.append(f"Autor: {author}")
        text.append(f"Erstellungsdatum: {datetime.date.today().strftime('%d.%m.%Y')}")
        text.append("")

        text.append(f"T_links [°C]: {project.get('T_left')}")
        text.append(f"T_∞ [°C]: {project.get('T_inf')}")
        text.append(f"h [W/m²K]: {project.get('h')}")
        text.append("")

        result = project.get("result", {})
        if result:
            text.append(f"Wärmestromdichte q = {result.get('q', 0):.3f} W/m²")
            text.append(f"Gesamtwiderstand = {result.get('R_total', 0):.5f} m²K/W")
        else:
            text.append("Keine Berechnungsergebnisse verfügbar.")

        if comment:
            text.append("\nKommentar:")
            text.append(comment)

        return "\n".join(text)
    except Exception as e:
        raise RuntimeError(f"Fehler beim Erstellen des Berichtinhalts: {e}")


def export_to_pdf(file_path: str, report_text: str, project_name: str,
                  project_data=None, image_path: str = None):
    """Exportiert Bericht als PDF mit tabellarischen Ergebnissen."""
    try:
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(file_path, pagesize=A4,
                                leftMargin=2 * cm, rightMargin=2 * cm)
        story = []
        story.append(Paragraph("<b>Heatrix – Isolierungsberechnung Bericht</b>",
                               styles["Title"]))
        story.append(Paragraph(f"<b>Projekt: {project_name}</b>",
                               styles["Heading2"]))
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(report_text.replace("\n", "<br/>"),
                               styles["Normal"]))
        story.append(Spacer(1, 0.5 * cm))

        if project_data and "result" in project_data:
            result = project_data["result"]
            T_if = result.get("interface_temperatures", [])
            T_avg = result.get("T_avg", [])
            k_avg = result.get("k_final", [])
            isolierungen = project_data.get("isolierungen", [])
            thicknesses = project_data.get("thicknesses", [])

            if T_if and T_avg and k_avg:
                data = [["Schicht", "Material", "Dicke [mm]",
                         "T_links [°C]", "T_rechts [°C]",
                         "T_mittel [°C]", "k_mittel [W/mK]"]]
                for i in range(len(T_if) - 1):
                    iso_name = isolierungen[i] if i < len(isolierungen) else f"Schicht {i+1}"
                    thick = thicknesses[i] if i < len(thicknesses) else 0.0
                    data.append([
                        f"{i+1}",
                        iso_name,
                        f"{thick:.2f}",
                        f"{T_if[i]:.2f}",
                        f"{T_if[i+1]:.2f}",
                        f"{T_avg[i]:.2f}",
                        f"{k_avg[i]:.4f}",
                    ])

                table = Table(data, hAlign='LEFT',
                              colWidths=[40, 110, 60, 65, 70, 70, 85])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d3d3d3")),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('ALIGN', (3, 1), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, -1), 'HeiseiMin-W3'),
                ]))
                story.append(table)
                story.append(Spacer(1, 0.5 * cm))

        if image_path and os.path.exists(image_path):
            story.append(Image(image_path, width=12 * cm, height=8 * cm))

        try:
            doc.build(story)
        except PermissionError:
            alt_path = file_path.replace(".pdf", "_neu.pdf")
            doc = SimpleDocTemplate(alt_path, pagesize=A4,
                                    leftMargin=2 * cm, rightMargin=2 * cm)
            doc.build(story)
        return True
    except Exception as e:
        raise RuntimeError(f"Fehler beim PDF-Export: {e}")


def export_to_docx(file_path: str, report_text: str, project_name: str,
                   project_data=None, image_path: str = None):
    """Exportiert Bericht als Word-Datei."""
    try:
        doc = Document()
        doc.add_heading("Heatrix – Isolierungsberechnung Bericht", level=1)
        doc.add_paragraph(f"Projekt: {project_name}")
        doc.add_paragraph(report_text)

        if project_data and "result" in project_data:
            result = project_data["result"]
            T_if = result.get("interface_temperatures", [])
            T_avg = result.get("T_avg", [])
            k_avg = result.get("k_final", [])
            isolierungen = project_data.get("isolierungen", [])
            thicknesses = project_data.get("thicknesses", [])

            if T_if and T_avg and k_avg:
                table = doc.add_table(rows=1, cols=7)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                hdr_cells = table.rows[0].cells
                headers = ["Schicht", "Material", "Dicke [mm]",
                           "T_links [°C]", "T_rechts [°C]",
                           "T_mittel [°C]", "k_mittel [W/mK]"]
                for cell, header in zip(hdr_cells, headers):
                    cell.text = header

                for i in range(len(T_if) - 1):
                    iso_name = isolierungen[i] if i < len(isolierungen) else f"Schicht {i+1}"
                    thick = thicknesses[i] if i < len(thicknesses) else 0.0
                    row = table.add_row().cells
                    row[0].text = str(i + 1)
                    row[1].text = iso_name
                    row[2].text = f"{thick:.2f}"
                    row[3].text = f"{T_if[i]:.2f}"
                    row[4].text = f"{T_if[i+1]:.3f}"
                    row[5].text = f"{T_avg[i]:.3f}"
                    row[6].text = f"{k_avg[i]:.4f}"

        if image_path and os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(6))

        try:
            doc.save(file_path)
        except PermissionError:
            alt_path = file_path.replace(".docx", "_neu.docx")
            doc.save(alt_path)
        return True
    except Exception as e:
        raise RuntimeError(f"Fehler beim Word-Export: {e}")